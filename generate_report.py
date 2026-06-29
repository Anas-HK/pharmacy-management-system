"""
generate_report.py  -  Build the project report as a Word (.docx) file.

Produces report/Pharmacy_Management_System_Report.docx containing the
project title and members, description, ERD, normalization, table
structures, DDL/DML and a section for GUI screenshots.

Run:
    python generate_report.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_DIR = os.path.join(BASE_DIR, "report")
OUT_FILE = os.path.join(OUT_DIR, "Pharmacy_Management_System_Report.docx")
ERD_FILE = os.path.join(OUT_DIR, "ERD.png")
SCHEMA_FILE = os.path.join(BASE_DIR, "schema.sql")

ACCENT = RGBColor(0x0b, 0x5d, 0x8a)


# --------------------------------------------------------------------- helpers
def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_code_box(doc, code, size=8.5):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F4F4")
    first = True
    for line in code.rstrip("\n").split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(size)


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


# --------------------------------------------------------------------- content
DESCRIPTION = (
    "The Pharmacy Management System is a desktop database application that helps a "
    "pharmacy manage its day to day operations. It keeps records of medicines and "
    "their stock levels, the categories medicines belong to, the suppliers who provide "
    "them, the customers who buy them, and the employees who serve those customers. It "
    "also records every sale and the individual medicines sold on each sale.\n\n"
    "The application is built with Python (Tkinter) for the graphical user interface "
    "and SQLite as the relational database. Through the GUI the user can add, update, "
    "delete and search records in every table, and can view a set of reports such as "
    "current stock, low stock alerts, medicines that are about to expire, sales per "
    "customer, revenue per category and full sales invoices.\n\n"
    "The database consists of seven related tables and is normalized up to the third "
    "normal form (3NF). Primary keys uniquely identify each record and foreign keys "
    "maintain the relationships and referential integrity between the tables."
)

OBJECTIVES = [
    "Design a normalized relational database (up to 3NF) for a pharmacy.",
    "Use appropriate primary keys, foreign keys, data types and constraints.",
    "Provide a GUI to Add, Update, Delete, Search and View records.",
    "Provide reports built from JOIN and aggregate (GROUP BY) queries.",
    "Demonstrate DDL (table creation) and DML (insert, update, delete, select).",
]

TOOLS = [
    ("Programming language", "Python 3"),
    ("GUI library", "Tkinter (ttk) - part of the Python standard library"),
    ("Database", "SQLite (file based relational database, standard SQL)"),
    ("Diagram", "Entity Relationship Diagram rendered with Matplotlib"),
]

NORMALIZATION = [
    ("Unnormalized data (the problem)",
     "If we stored a whole sale in a single row, for example one Sale row that lists all "
     "the medicines bought as a comma separated value, and repeated the customer name, "
     "phone and address on every sale, the data would contain repeating groups and a lot "
     "of duplication. This causes insertion, update and deletion anomalies. Normalization "
     "removes these problems by splitting the data into well defined tables."),
    ("First Normal Form (1NF) - atomic values, no repeating groups",
     "Every column holds a single (atomic) value and every table has a primary key. A sale "
     "that contains several medicines is not stored as a list inside the Sale table. "
     "Instead, each medicine on a sale is stored as a separate row in the Sale_Item table. "
     "So no cell contains multiple values and there are no repeating groups."),
    ("Second Normal Form (2NF) - no partial dependency",
     "The database is in 1NF and every non-key column depends on the whole primary key. "
     "Each table uses a single-column primary key (for example medicine_id, sale_id, "
     "sale_item_id), so no column can depend on only part of the key. For instance, in "
     "Sale_Item the quantity and unit_price depend on the full key sale_item_id, which "
     "identifies one specific line of one specific sale."),
    ("Third Normal Form (3NF) - no transitive dependency",
     "The database is in 2NF and no non-key column depends on another non-key column. For "
     "example, the Medicine table stores only category_id and supplier_id (foreign keys); "
     "it does not store category_name or supplier_name. Those descriptive attributes live "
     "once in the Category and Supplier tables. This removes the transitive dependency "
     "medicine_id -> category_id -> category_name. Likewise, the Sale table stores "
     "customer_id rather than the customer's name, phone and address, which are kept once "
     "in the Customer table."),
    ("Result",
     "Each fact is stored in exactly one place, update and deletion anomalies are avoided, "
     "and referential integrity is enforced through foreign keys."),
]

RELATIONSHIPS = [
    "One Category has many Medicines.  (Category 1 : N Medicine)",
    "One Supplier supplies many Medicines.  (Supplier 1 : N Medicine)",
    "One Customer can make many Sales.  (Customer 1 : N Sale)",
    "One Employee can handle many Sales.  (Employee 1 : N Sale)",
    "One Sale contains many Sale_Items.  (Sale 1 : N Sale_Item)",
    "One Medicine can appear in many Sale_Items.  (Medicine 1 : N Sale_Item)",
    "Because of the Sale_Item table, Sale and Medicine have a many-to-many (M : N) "
    "relationship: a sale can include many medicines and a medicine can be sold on many "
    "sales. Sale_Item resolves this M:N relationship.",
]

# table -> (purpose, [(column, datatype, key/constraint, description)])
TABLE_DOCS = {
    "Category": ("Groups of medicines (for example Antibiotics).", [
        ("category_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the category"),
        ("category_name", "TEXT", "NOT NULL, UNIQUE", "Name of the category"),
        ("description", "TEXT", "(nullable)", "Short description"),
    ]),
    "Supplier": ("Companies that supply medicines to the pharmacy.", [
        ("supplier_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the supplier"),
        ("supplier_name", "TEXT", "NOT NULL", "Company name"),
        ("contact_person", "TEXT", "(nullable)", "Person to contact"),
        ("phone", "TEXT", "(nullable)", "Phone number"),
        ("email", "TEXT", "(nullable)", "Email address"),
        ("address", "TEXT", "(nullable)", "Address"),
    ]),
    "Medicine": ("Products in stock; each has one category and one supplier.", [
        ("medicine_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the medicine"),
        ("medicine_name", "TEXT", "NOT NULL", "Name of the medicine"),
        ("category_id", "INTEGER", "NOT NULL, FOREIGN KEY -> Category", "Category it belongs to"),
        ("supplier_id", "INTEGER", "NOT NULL, FOREIGN KEY -> Supplier", "Supplier that provides it"),
        ("unit_price", "REAL", "NOT NULL, CHECK >= 0", "Selling price per unit"),
        ("stock_quantity", "INTEGER", "NOT NULL, DEFAULT 0, CHECK >= 0", "Units currently in stock"),
        ("expiry_date", "TEXT", "(nullable)", "Expiry date (YYYY-MM-DD)"),
        ("batch_no", "TEXT", "(nullable)", "Manufacturer batch number"),
    ]),
    "Customer": ("People who buy medicines.", [
        ("customer_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the customer"),
        ("customer_name", "TEXT", "NOT NULL", "Customer name"),
        ("phone", "TEXT", "(nullable)", "Phone number"),
        ("email", "TEXT", "(nullable)", "Email address"),
        ("address", "TEXT", "(nullable)", "Address"),
    ]),
    "Employee": ("Pharmacy staff who handle the sales.", [
        ("employee_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the employee"),
        ("employee_name", "TEXT", "NOT NULL", "Employee name"),
        ("role", "TEXT", "(nullable)", "Job role (Pharmacist, Cashier, ...)"),
        ("phone", "TEXT", "(nullable)", "Phone number"),
        ("salary", "REAL", "CHECK >= 0", "Monthly salary"),
        ("hire_date", "TEXT", "(nullable)", "Date of joining"),
    ]),
    "Sale": ("One sales transaction (the invoice header).", [
        ("sale_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the sale"),
        ("customer_id", "INTEGER", "NOT NULL, FOREIGN KEY -> Customer", "Customer who bought"),
        ("employee_id", "INTEGER", "NOT NULL, FOREIGN KEY -> Employee", "Employee who served"),
        ("sale_date", "TEXT", "NOT NULL", "Date of the sale"),
        ("payment_method", "TEXT", "(nullable)", "Cash / Card / Mobile Banking"),
        ("total_amount", "REAL", "DEFAULT 0", "Total value (sum of its line items)"),
    ]),
    "Sale_Item": ("Line items of a sale (the invoice detail). Resolves the M:N "
                  "relationship between Sale and Medicine.", [
        ("sale_item_id", "INTEGER", "PRIMARY KEY, AUTOINCREMENT", "Unique id of the line item"),
        ("sale_id", "INTEGER", "NOT NULL, FOREIGN KEY -> Sale (ON DELETE CASCADE)", "Sale it belongs to"),
        ("medicine_id", "INTEGER", "NOT NULL, FOREIGN KEY -> Medicine", "Medicine that was sold"),
        ("quantity", "INTEGER", "NOT NULL, CHECK > 0", "Number of units sold"),
        ("unit_price", "REAL", "NOT NULL, CHECK >= 0", "Price per unit at time of sale"),
        ("subtotal", "REAL", "(computed)", "quantity x unit_price"),
    ]),
}

DML_SAMPLES = [
    ("Insert a new medicine (DML - INSERT)",
     "INSERT INTO Medicine (medicine_name, category_id, supplier_id,\n"
     "                      unit_price, stock_quantity, expiry_date, batch_no)\n"
     "VALUES ('Azithromycin 250mg', 1, 3, 9.50, 120, '2027-10-31', 'BATCH-AZI-21');"),
    ("Update the stock of a medicine (DML - UPDATE)",
     "UPDATE Medicine\n"
     "SET stock_quantity = stock_quantity + 100\n"
     "WHERE medicine_name = 'Paracetamol 500mg';"),
    ("Delete a record (DML - DELETE)",
     "DELETE FROM Customer\n"
     "WHERE customer_id = 20;"),
    ("Search a medicine by partial name (SELECT + LIKE)",
     "SELECT * FROM Medicine\n"
     "WHERE medicine_name LIKE '%cin%';"),
    ("Join medicines with their category and supplier (INNER JOIN)",
     "SELECT m.medicine_name, c.category_name, s.supplier_name, m.unit_price\n"
     "FROM Medicine m\n"
     "JOIN Category c ON m.category_id = c.category_id\n"
     "JOIN Supplier s ON m.supplier_id = s.supplier_id;"),
    ("Total amount spent by each customer (GROUP BY + aggregate)",
     "SELECT cu.customer_name, COUNT(sl.sale_id) AS orders,\n"
     "       SUM(sl.total_amount) AS total_spent\n"
     "FROM Customer cu\n"
     "JOIN Sale sl ON cu.customer_id = sl.customer_id\n"
     "GROUP BY cu.customer_id\n"
     "ORDER BY total_spent DESC;"),
    ("Medicines priced above the average (subquery)",
     "SELECT medicine_name, unit_price\n"
     "FROM Medicine\n"
     "WHERE unit_price > (SELECT AVG(unit_price) FROM Medicine);"),
]

SCREENSHOTS = [
    ("1_main_window_view.png",
     "Main window: the Medicine tab showing all records (the View function)."),
    ("2_add_record.png",
     "Add: the Category form filled in, ready to click Add."),
    ("3_update_record.png",
     "Update: a record selected from the list with a field edited, ready to click Update."),
    ("4_delete_confirmation.png",
     "Delete: the confirmation dialog shown before a record is removed."),
    ("5_search.png",
     "Search: medicines filtered by the search term 'Amox'."),
    ("6_reports.png",
     "Reports: the Invoice details report, built from a multi-table join."),
]


# --------------------------------------------------------------------- builder
def build():
    os.makedirs(OUT_DIR, exist_ok=True)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title page ----
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pharmacy Management System")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Database Management System Lab Project")
    run.font.size = Pt(15)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by")
    run.bold = True
    run.font.size = Pt(12)

    members = doc.add_table(rows=1, cols=3)
    members.style = "Table Grid"
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = members.rows[0].cells
    for cell, text in zip(hdr, ("Name", "Student ID", "Role")):
        cell.paragraphs[0].add_run(text).bold = True
        shade_cell(cell, "D9E6EF")
    for i in range(1, 4):
        row = members.add_row().cells
        row[0].text = f"[Member {i} Full Name]"
        row[1].text = "[Student ID]"
        row[2].text = "[Role]"

    doc.add_paragraph()
    for label in ("Course: [Course Code - Course Title]",
                  "Submitted to: [Instructor / Teacher Name]",
                  "Department / Section: [Your Department and Section]",
                  "Date of submission: [Date]"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label).font.size = Pt(12)

    doc.add_page_break()

    # ---- 1. Introduction / Objectives ----
    doc.add_heading("1. Project Description", level=1)
    for para in DESCRIPTION.split("\n\n"):
        doc.add_paragraph(para)

    doc.add_heading("Objectives", level=2)
    for obj in OBJECTIVES:
        doc.add_paragraph(obj, style="List Bullet")

    doc.add_heading("Tools and Technologies", level=2)
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    for cell, text in zip(t.rows[0].cells, ("Item", "Technology used")):
        cell.paragraphs[0].add_run(text).bold = True
        shade_cell(cell, "D9E6EF")
    for item, tech in TOOLS:
        row = t.add_row().cells
        row[0].text = item
        row[1].text = tech

    # ---- 2. ERD ----
    doc.add_heading("2. Entity Relationship Diagram (ERD)", level=1)
    doc.add_paragraph(
        "The ERD below shows the seven entities, their attributes, the primary key (PK) "
        "and foreign key (FK) of each entity, and the one-to-many (1:N) relationships "
        "between them.")
    if os.path.exists(ERD_FILE):
        doc.add_picture(ERD_FILE, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Figure 1: ER Diagram of the Pharmacy Management System")

    doc.add_heading("Relationships", level=2)
    for rel in RELATIONSHIPS:
        doc.add_paragraph(rel, style="List Bullet")

    # ---- 3. Normalization ----
    doc.add_page_break()
    doc.add_heading("3. Normalization (up to 3NF)", level=1)
    for title, text in NORMALIZATION:
        doc.add_heading(title, level=2)
        doc.add_paragraph(text)

    # ---- 4. Table structure ----
    doc.add_page_break()
    doc.add_heading("4. Table Structure", level=1)
    doc.add_paragraph(
        "The database has seven tables. Each table is listed below with its columns, "
        "data types, keys and constraints.")
    for name, (purpose, columns) in TABLE_DOCS.items():
        doc.add_heading(name, level=2)
        doc.add_paragraph(purpose)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ("Column", "Data Type", "Key / Constraint", "Description")
        for cell, text in zip(table.rows[0].cells, headers):
            cell.paragraphs[0].add_run(text).bold = True
            shade_cell(cell, "D9E6EF")
        for col, dtype, key, desc in columns:
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = col, dtype, key, desc

    # ---- 5. DDL ----
    doc.add_page_break()
    doc.add_heading("5. DDL - Data Definition Language", level=1)
    doc.add_paragraph(
        "The following SQL statements create the database tables with their primary keys, "
        "foreign keys, data types and constraints. (Full file: schema.sql)")
    with open(SCHEMA_FILE, "r", encoding="utf-8") as f:
        schema_sql = f.read()
    add_code_box(doc, schema_sql)

    # ---- 6. DML ----
    doc.add_page_break()
    doc.add_heading("6. DML - Data Manipulation Language", level=1)
    doc.add_paragraph(
        "At least 20 records are inserted into every table (see seed.sql). A complete set "
        "of 20 demonstration queries is provided in queries.sql. Representative examples "
        "are shown below.")
    for title, sql in DML_SAMPLES:
        doc.add_heading(title, level=2)
        add_code_box(doc, sql)

    # ---- 7. GUI screenshots ----
    doc.add_page_break()
    doc.add_heading("7. GUI Screenshots", level=1)
    doc.add_paragraph(
        "The screenshots below show the running application performing each of the "
        "required operations: viewing, adding, updating, deleting, searching, and "
        "viewing reports.")
    shot_dir = os.path.join(OUT_DIR, "screenshots")
    for i, (fname, caption) in enumerate(SCREENSHOTS, start=1):
        p = doc.add_paragraph()
        p.add_run(f"Screenshot {i}: {caption}").bold = True
        path = os.path.join(shot_dir, fname)
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            ph = doc.add_paragraph()
            run = ph.add_run("[ Paste screenshot here ]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.add_paragraph()

    # ---- 8. How to run ----
    doc.add_heading("8. How to Run the Project", level=1)
    steps = [
        "Make sure Python 3 is installed (the GUI uses tkinter and sqlite3, which come "
        "with Python, so nothing extra is needed to run the app).",
        "Open a terminal in the project folder.",
        "Build the database (creates pharmacy.db with all the sample data):   python init_db.py",
        "Start the application:   python app.py    (the database is also built "
        "automatically the first time if it is missing).",
        "Use the tabs to Add, Update, Delete and Search records, and open the Reports tab "
        "to view the reports.",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")

    doc.save(OUT_FILE)
    print("Report written to:", OUT_FILE)


if __name__ == "__main__":
    build()
